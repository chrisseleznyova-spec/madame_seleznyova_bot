import os
import logging
import io
import re
import asyncio
import json
from datetime import datetime, date
from aiogram import Bot, Dispatcher, types, F
from aiogram.filters import CommandStart, Command
from aiogram.fsm.context import FSMContext
from aiogram.fsm.state import State, StatesGroup
from aiogram.fsm.storage.memory import MemoryStorage
from aiogram.types import (
    ReplyKeyboardMarkup, KeyboardButton, ReplyKeyboardRemove,
    InlineKeyboardMarkup, InlineKeyboardButton, BufferedInputFile
)
import anthropic
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import asyncpg

logging.basicConfig(level=logging.INFO)

BOT_TOKEN = os.environ.get("BOT_TOKEN")
ANTHROPIC_API_KEY = os.environ.get("ANTHROPIC_API_KEY")
CHANNEL_URL = "https://t.me/seleznyovaochemzadymalas"
SESSION_URL = os.environ.get("SESSION_URL", "https://t.me/")
ADMIN_ID = int(os.environ.get("ADMIN_ID", "0"))
DATABASE_URL = os.environ.get("DATABASE_PUBLIC_URL")

bot = Bot(token=BOT_TOKEN, parse_mode="HTML")
dp = Dispatcher(storage=MemoryStorage())
anthropic_client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)

# --- PostgreSQL ---

async def get_db():
    return await asyncpg.connect(DATABASE_URL)

async def init_db():
    conn = await get_db()
    await conn.execute("""
        CREATE TABLE IF NOT EXISTS users (
            user_id BIGINT PRIMARY KEY,
            username TEXT,
            sessions INT DEFAULT 0,
            completed INT DEFAULT 0,
            first_seen TIMESTAMP DEFAULT NOW(),
            last_seen TIMESTAMP DEFAULT NOW(),
            source TEXT DEFAULT '',
            sphere TEXT DEFAULT '',
            last_theme TEXT DEFAULT '',
            remind_at TIMESTAMP,
            reminded BOOLEAN DEFAULT FALSE
        );
        CREATE TABLE IF NOT EXISTS themes (
            id SERIAL PRIMARY KEY,
            user_id BIGINT,
            theme TEXT,
            created_at DATE DEFAULT CURRENT_DATE
        );
        CREATE TABLE IF NOT EXISTS spheres (
            id SERIAL PRIMARY KEY,
            user_id BIGINT,
            sphere TEXT,
            created_at DATE DEFAULT CURRENT_DATE
        );
        CREATE TABLE IF NOT EXISTS feedback (
            id SERIAL PRIMARY KEY,
            user_id BIGINT,
            value TEXT,
            created_at DATE DEFAULT CURRENT_DATE
        );
    """)
    await conn.close()

async def record_session(user_id: int, username: str, source: str = ""):
    conn = await get_db()
    try:
        existing = await conn.fetchrow("SELECT user_id FROM users WHERE user_id = $1", user_id)
        if existing:
            await conn.execute("""
                UPDATE users SET sessions = sessions + 1, last_seen = NOW(), username = $2
                WHERE user_id = $1
            """, user_id, username)
        else:
            await conn.execute("""
                INSERT INTO users (user_id, username, sessions, source)
                VALUES ($1, $2, 1, $3)
            """, user_id, username, source)
    finally:
        await conn.close()

async def record_theme(user_id: int, theme: str):
    conn = await get_db()
    try:
        await conn.execute("INSERT INTO themes (user_id, theme) VALUES ($1, $2)", user_id, theme)
        await conn.execute("UPDATE users SET last_theme = $2 WHERE user_id = $1", user_id, theme)
    finally:
        await conn.close()

async def record_completion(user_id: int):
    conn = await get_db()
    try:
        await conn.execute("UPDATE users SET completed = completed + 1 WHERE user_id = $1", user_id)
    finally:
        await conn.close()

async def record_sphere(user_id: int, sphere: str):
    conn = await get_db()
    try:
        await conn.execute("UPDATE users SET sphere = $2 WHERE user_id = $1", user_id, sphere)
        await conn.execute("INSERT INTO spheres (user_id, sphere) VALUES ($1, $2)", user_id, sphere)
    finally:
        await conn.close()

async def record_feedback(user_id: int, value: str):
    conn = await get_db()
    try:
        await conn.execute("INSERT INTO feedback (user_id, value) VALUES ($1, $2)", user_id, value)
    finally:
        await conn.close()

async def set_reminder(user_id: int, theme: str):
    from datetime import timedelta
    remind_dt = datetime.now() + timedelta(days=3)
    conn = await get_db()
    try:
        await conn.execute("""
            UPDATE users SET last_theme = $2, remind_at = $3, reminded = FALSE
            WHERE user_id = $1
        """, user_id, theme, remind_dt)
    finally:
        await conn.close()

async def get_stats():
    conn = await get_db()
    try:
        total_users = await conn.fetchval("SELECT COUNT(*) FROM users")
        total_sessions = await conn.fetchval("SELECT COALESCE(SUM(sessions), 0) FROM users")
        total_completed = await conn.fetchval("SELECT COALESCE(SUM(completed), 0) FROM users")
        today_sessions = await conn.fetchval(
            "SELECT COUNT(*) FROM users WHERE DATE(last_seen) = CURRENT_DATE"
        )
        top_themes = await conn.fetch(
            "SELECT theme, COUNT(*) as cnt FROM themes GROUP BY theme ORDER BY cnt DESC LIMIT 8"
        )
        feedback_rows = await conn.fetch(
            "SELECT value, COUNT(*) as cnt FROM feedback GROUP BY value"
        )
        top_spheres = await conn.fetch(
            "SELECT sphere, COUNT(*) as cnt FROM spheres GROUP BY sphere ORDER BY cnt DESC LIMIT 4"
        )
        returning = await conn.fetch(
            "SELECT user_id, username, sessions FROM users WHERE sessions > 1 ORDER BY sessions DESC LIMIT 5"
        )
        all_themes = await conn.fetch("SELECT theme FROM themes ORDER BY created_at DESC LIMIT 50")
        return {
            "total_users": total_users,
            "total_sessions": total_sessions,
            "total_completed": total_completed,
            "today_sessions": today_sessions,
            "top_themes": top_themes,
            "feedback": feedback_rows,
            "top_spheres": top_spheres,
            "returning": returning,
            "all_themes": [r["theme"] for r in all_themes],
        }
    finally:
        await conn.close()

# --- Системный промт ---

SYSTEM_PROMPT = """Ты — ИИ-ассистент Кристины Селезнёвой, коуча ICF.
Бот работает как персонаж «Мадам Селезнёва» и помогает людям коротко разобрать их жизненную ситуацию.
Твоя задача — через несколько точных вопросов собрать картину происходящего и дать короткий, ясный разбор.

СТИЛЬ ОБЩЕНИЯ:
- коротко и по делу
- без воды
- поддерживающе, но прямо
- как умный внимательный собеседник, а не психолог из сериала
- тон — спокойный, немного ироничный

Никогда не используй слова: «однозначно», «безусловно», «конечно».
Не давай банальных советов.

ФОРМАТ ДИАЛОГА:
- один вопрос = одно сообщение
- не задавай несколько вопросов одновременно
- вопросы: 1-2 предложения

УТОЧНЯЮЩИЕ ВОПРОСЫ:
Задавай 4-6 вопросов. Если после 6 картина ясна — переходи к разбору.
Если есть 70-80% понимания — переходи к разбору.

КОГДА ПЕРЕХОДИШЬ К РАЗБОРУ — НЕ пиши никакой вводной фразы типа «Я внимательно посмотрю на ваши ответы». Сразу начинай с первого раздела разбора.

СТРУКТУРА ФИНАЛЬНОГО РАЗБОРА (до 1500 символов):
СТРОГО используй HTML-теги <b>...</b> для заголовков разделов. Никаких звёздочек **текст** — только <b>текст</b>.

<b>Что происходит</b>
[текст]

<b>Почему это происходит</b>
[текст]

<b>Эмоциональный маркер</b>
[текст]

<b>На что стоит обратить внимание</b>
[текст]

Затем:
Оставлю вам один вопрос, который может быть полезно спокойно обдумать:

[вопрос]

Иногда именно с него начинает постепенно распутываться вся ситуация.

<i>Этот вопрос — не для ответа в чате, а для себя.</i>

<i>Этот разбор сделан в боте «Мадам Селезнёва разбирает» — проекте Кристины Селезнёвой.</i>"""


class Dialog(StatesGroup):
    consent = State()
    mini_sphere = State()
    start = State()
    describe = State()
    questions = State()
    final = State()
    post_final = State()


def btn(labels: list[str]) -> ReplyKeyboardMarkup:
    return ReplyKeyboardMarkup(
        keyboard=[[KeyboardButton(text=t)] for t in labels],
        resize_keyboard=True,
        one_time_keyboard=True
    )


async def ask_claude(messages: list[dict]) -> str:
    response = anthropic_client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=1500,
        system=SYSTEM_PROMPT,
        messages=messages
    )
    return response.content[0].text


def html_to_plain(text: str) -> str:
    text = re.sub(r'<b>(.*?)</b>', r'\1', text)
    text = re.sub(r'<i>(.*?)</i>', r'\1', text)
    text = re.sub(r'<[^>]+>', '', text)
    return text


def create_docx(final_text: str) -> bytes:
    doc = Document()
    title = doc.add_heading('Мадам Селезнёва разбирает', level=1)
    if title.runs:
        title.runs[0].font.color.rgb = RGBColor(0x6B, 0x3F, 0xA0)
    sub = doc.add_paragraph('Разбор вашей ситуации')
    if sub.runs:
        sub.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        sub.runs[0].font.size = Pt(11)
    doc.add_paragraph()
    text = re.sub(r'<b>(.*?)</b>', r'\n▶ \1\n', final_text)
    text = re.sub(r'<i>(.*?)</i>', r'\1', text)
    text = re.sub(r'<[^>]+>', '', text)
    for line in text.split('\n'):
        line = line.strip()
        if not line:
            continue
        if line.startswith('▶ '):
            heading_text = line[2:]
            h = doc.add_heading(heading_text, level=2)
            if h.runs:
                h.runs[0].font.color.rgb = RGBColor(0x3D, 0x20, 0x60)
        else:
            p = doc.add_paragraph(line)
            if p.runs:
                p.runs[0].font.size = Pt(11)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


WELCOME_PHOTO = os.environ.get("WELCOME_PHOTO_ID", "")
WELCOME_PHOTO_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "ChatGPT Image 14 мар. 2026 г., 21_30_47.png")
PRIVACY_URL = "https://telegra.ph/Politika-konfidencialnosti-03-13-46"


@dp.message(CommandStart())
async def cmd_start(message: types.Message, state: FSMContext):
    try:
        await state.clear()
        await state.set_state(Dialog.consent)
        user = message.from_user
        args = message.text.split()
        source = args[1] if len(args) > 1 else "direct"
        try:
            await record_session(user.id, user.username or user.full_name or "", source=source)
        except Exception as e:
            logging.error(f"record_session error: {e}")
        await message.answer(
            "Прежде чем начать — один момент.\n\n"
            "В ходе разбора вы будете делиться личными переживаниями. "
            "Ваши ответы не сохраняются и не передаются третьим лицам после завершения сессии.\n\n"
            "Нажимая «Принимаю условия», вы подтверждаете согласие на обработку текста ваших сообщений в рамках этого разбора.",
            reply_markup=InlineKeyboardMarkup(inline_keyboard=[
                [InlineKeyboardButton(text="📄 Политика конфиденциальности", url=PRIVACY_URL)],
                [InlineKeyboardButton(text="✅ Принимаю условия", callback_data="consent_accept")],
            ])
        )
    except Exception as e:
        logging.error(f"cmd_start error: {e}")
        await message.answer("Что-то пошло не так. Попробуйте ещё раз — /start")


@dp.callback_query(F.data == "consent_accept")
async def consent_handler(callback: types.CallbackQuery, state: FSMContext):
    await callback.answer()
    await state.set_state(Dialog.mini_sphere)
    caption = (
        "Здравствуйте.\n"
        "Я — Мадам Селезнёва.\n\n"
        "Задам несколько точных вопросов и попробую собрать картину вашей ситуации.\n\n"
        "Это не терапия и не диагноз.\n"
        "Но иногда уже по ответам становится видно, где именно всё запуталось."
    )
    if WELCOME_PHOTO:
        await callback.message.answer_photo(photo=WELCOME_PHOTO, caption=caption)
    elif os.path.exists(WELCOME_PHOTO_PATH):
        photo_file = types.FSInputFile(WELCOME_PHOTO_PATH)
        await callback.message.answer_photo(photo=photo_file, caption=caption)
    else:
        await callback.message.answer(caption)
    await asyncio.sleep(1)
    await callback.message.answer(
        "Один вопрос перед началом — чтобы я лучше понимала контекст.\n\n"
        "Какая сфера жизни сейчас занимает больше всего мыслей?\n"
        "<i>Выберите до двух вариантов, затем нажмите «Готово»</i>",
        reply_markup=sphere_keyboard([])
    )


def sphere_keyboard(selected: list) -> InlineKeyboardMarkup:
    options = ["Отношения", "Работа и карьера", "Семья", "Я сама / внутреннее состояние"]
    buttons = []
    for opt in options:
        mark = "✅ " if opt in selected else ""
        buttons.append([InlineKeyboardButton(text=f"{mark}{opt}", callback_data=f"sphere_{opt}")])
    buttons.append([InlineKeyboardButton(text="Готово →", callback_data="sphere_done")])
    return InlineKeyboardMarkup(inline_keyboard=buttons)


@dp.callback_query(Dialog.mini_sphere, F.data.startswith("sphere_") & ~F.data.endswith("done"))
async def sphere_toggle(callback: types.CallbackQuery, state: FSMContext):
    chosen = callback.data[7:]
    data = await state.get_data()
    selected = data.get("selected_spheres", [])
    if chosen in selected:
        selected.remove(chosen)
    elif len(selected) < 2:
        selected.append(chosen)
    await state.update_data(selected_spheres=selected)
    await callback.message.edit_reply_markup(reply_markup=sphere_keyboard(selected))
    await callback.answer()


@dp.callback_query(Dialog.mini_sphere, F.data == "sphere_done")
async def sphere_done(callback: types.CallbackQuery, state: FSMContext):
    await callback.answer()
    data = await state.get_data()
    selected = data.get("selected_spheres", [])
    if not selected:
        await callback.answer("Выберите хотя бы один вариант", show_alert=True)
        return
    sphere_str = ", ".join(selected)
    try:
        await record_sphere(callback.from_user.id, sphere_str)
    except Exception as e:
        logging.error(f"record_sphere error: {e}")
    await state.set_state(Dialog.start)
    await callback.message.edit_reply_markup(reply_markup=None)
    await callback.message.answer("Хорошо. Если готовы — начнём.", reply_markup=btn(["Разобрать ситуацию"]))


@dp.message(Dialog.start, F.text == "Разобрать ситуацию")
async def screen_describe(message: types.Message, state: FSMContext):
    await state.set_state(Dialog.describe)
    await message.answer(
        "Расскажите немного о том, что сейчас происходит в вашей жизни.\n"
        "Что вас беспокоит или с чем вы хотели бы разобраться?\n\n"
        "Можно коротко — до 500 символов.",
        reply_markup=ReplyKeyboardRemove()
    )


@dp.message(Dialog.describe)
async def screen_themes(message: types.Message, state: FSMContext):
    user_text = message.text[:500]
    history = [{"role": "user", "content": f"Моя ситуация: {user_text}"}]
    await state.update_data(situation=user_text, history=history)
    await message.answer("Анализирую...", reply_markup=ReplyKeyboardRemove())
    history.append({
        "role": "user",
        "content": (
            "Назови 2-3 темы для разбора — каждая на отдельной строке. "
            "Формат: короткая назывная фраза, без вопросов, без нумерации, до 50 символов. "
            "Пример: 'Страх после пережитого', 'Возвращение к нормальной жизни'. "
            "Если видишь только 2 реальные темы — дай 2, не выдумывай третью. "
            "Только темы, ничего лишнего."
        )
    })
    response = await ask_claude(history)
    history.append({"role": "assistant", "content": response})
    lines = [l.strip() for l in response.strip().split("\n") if l.strip()]
    skip_words = ["вот", "варианта", "варианты", "предлагаю", "можно разобрать", "три", "рассмотрим", "тем"]
    themes = []
    for l in lines:
        l_lower = l.lower()
        if any(w in l_lower for w in skip_words):
            continue
        if len(l) < 80:
            themes.append(l)
        if len(themes) == 3:
            break
    if len(themes) < 2:
        themes = lines[:2] if len(lines) >= 2 else (lines + ["Другое"])[:2]
    numbered = "\n".join(f"{i+1}. {t}" for i, t in enumerate(themes))
    next_num = len(themes) + 1
    text = (
        f"Похоже, что здесь может идти речь о:\n{numbered}\n"
        f"{next_num}. Другое — можете уточнить своими словами\n\n"
        "Я правильно понимаю направление вашей ситуации?"
    )
    await state.update_data(history=history, themes=themes, question_count=0)
    await state.set_state(Dialog.questions)
    await message.answer(text, reply_markup=btn([*themes, "Другое"]))


@dp.message(Dialog.questions)
async def handle_questions(message: types.Message, state: FSMContext):
    data = await state.get_data()
    history = data.get("history", [])
    q_count = data.get("question_count", 0)
    user_input = message.text

    if q_count == 0:
        try:
            await record_theme(message.from_user.id, user_input)
        except Exception as e:
            logging.error(f"record_theme error: {e}")
        await message.answer(
            "Вопросы будут личными — это и есть суть разбора.\n"
            "Чем честнее ответите, тем точнее картина.\n\n"
            "<i>Всё конфиденциально: мы не собираем и не храним ваши ответы.</i>",
            reply_markup=ReplyKeyboardRemove()
        )
        history.append({"role": "user", "content": f"Пользователь выбрал направление: {user_input}. Задай первый уточняющий вопрос по этой теме — один вопрос, коротко, 1-2 предложения. Не повторяй формулировку темы."})
        question = await ask_claude(history)
        history.append({"role": "assistant", "content": question})
        await state.update_data(history=history, question_count=1)
        await message.answer(f"<b>Вопрос 1 из 4</b>\n\n{question}")
        return

    history.append({"role": "user", "content": user_input})

    if q_count >= 6:
        await state.update_data(history=history)
        await do_final(message, state)
        return

    history.append({
        "role": "user",
        "content": (
            f"Это был ответ на вопрос №{q_count}. "
            "Если уже есть 70-80% понимания — напиши ТОЛЬКО слово РАЗБОР. "
            "Если нужен ещё вопрос — задай один, коротко."
        )
    })
    response = await ask_claude(history)

    if "РАЗБОР" in response.upper() or q_count >= 5:
        await state.update_data(history=history)
        await do_final(message, state)
        return

    history.append({"role": "assistant", "content": response})
    await state.update_data(history=history, question_count=q_count + 1)
    progress = f"<b>Вопрос {q_count + 1} из 4</b>\n\n"
    await message.answer(f"{progress}{response}")


async def do_final(message: types.Message, state: FSMContext):
    data = await state.get_data()
    history = data.get("history", [])
    await message.answer("Собираю разбор...", reply_markup=ReplyKeyboardRemove())
    history.append({
        "role": "user",
        "content": (
            "Сделай финальный разбор. "
            "ВАЖНО: используй ТОЛЬКО HTML-теги <b>заголовок</b> для разделов. "
            "Никаких **звёздочек**, никаких вводных фраз типа 'Я внимательно посмотрю'.\n\n"
            "Начни сразу с:\n"
            "<b>Что происходит</b>\n[текст]\n\n"
            "<b>Почему это происходит</b>\n[текст]\n\n"
            "<b>Эмоциональный маркер</b>\n[текст]\n\n"
            "<b>На что стоит обратить внимание</b>\n[текст]\n\n"
            "Затем:\nОставлю вам один вопрос, который может быть полезно спокойно обдумать:\n\n<b>[вопрос]</b>\n\n"
            "Иногда именно с него начинает постепенно распутываться вся ситуация.\n\n"
            "<i>Этот разбор сделан в боте «Мадам Селезнёва разбирает» — проекте Кристины Селезнёвой.</i>\n\n"
            "Весь текст — до 1500 символов. Никаких --- разделителей."
        )
    })
    final_text = await ask_claude(history)
    history.append({"role": "assistant", "content": final_text})
    await state.update_data(history=history, final_text=final_text)
    await state.set_state(Dialog.final)

    try:
        await record_completion(message.from_user.id)
    except Exception as e:
        logging.error(f"record_completion error: {e}")

    data = await state.get_data()
    themes = data.get("themes", [])
    chosen_theme = themes[0] if themes else ""
    try:
        await set_reminder(message.from_user.id, chosen_theme)
    except Exception as e:
        logging.error(f"set_reminder error: {e}")

    keyboard = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="📎 Сохранить разбор (Word)", callback_data="save")],
    ])
    await message.answer(final_text, reply_markup=keyboard)


@dp.callback_query(F.data == "save")
async def save_handler(callback: types.CallbackQuery, state: FSMContext):
    data = await state.get_data()
    final_text = data.get("final_text", "")
    await callback.answer("Генерирую файл...")
    try:
        docx_bytes = create_docx(final_text)
        docx_file = BufferedInputFile(docx_bytes, filename="razбор_madame_seleznyova.docx")
        await callback.message.answer_document(docx_file, caption="Ваш разбор сохранён 📎")
    except Exception as e:
        logging.error(f"DOCX error: {e}")
        await callback.message.answer("Не удалось создать файл. Разбор сохранён выше в чате.")
    await after_final(callback.message, state)


async def after_final(message: types.Message, state: FSMContext):
    await state.set_state(Dialog.post_final)
    await asyncio.sleep(1)
    await message.answer(
        "Скажите — этот разбор был для вас полезным?",
        reply_markup=InlineKeyboardMarkup(inline_keyboard=[
            [
                InlineKeyboardButton(text="✅ Да, попал в точку", callback_data="fb_yes"),
                InlineKeyboardButton(text="🤔 Частично", callback_data="fb_partly"),
            ],
            [InlineKeyboardButton(text="❌ Не очень", callback_data="fb_no")],
        ])
    )


@dp.callback_query(F.data.startswith("fb_"))
async def feedback_handler(callback: types.CallbackQuery, state: FSMContext):
    await callback.answer()
    fb_map = {"fb_yes": "да", "fb_partly": "частично", "fb_no": "нет"}
    fb_value = fb_map.get(callback.data, "")
    try:
        await record_feedback(callback.from_user.id, fb_value)
    except Exception as e:
        logging.error(f"record_feedback error: {e}")

    if callback.data == "fb_yes":
        reply = "Рада слышать. Значит, картина сложилась."
    elif callback.data == "fb_partly":
        reply = "Понятно. Иногда одного разбора недостаточно — ситуация многослойная."
    else:
        reply = "Жаль. Возможно, тема оказалась сложнее или вопросы не попали в нужное место."

    await callback.message.answer(reply)
    await asyncio.sleep(1)
    await callback.message.answer(
        "Если вам близок такой способ разбираться в сложных ситуациях —\n"
        "в моём канале я регулярно публикую похожие разборы и наблюдения.",
        reply_markup=InlineKeyboardMarkup(inline_keyboard=[
            [InlineKeyboardButton(text="Перейти в канал", url=CHANNEL_URL)]
        ])
    )
    await asyncio.sleep(1)
    await callback.message.answer(
        "Иногда одного разбора достаточно, чтобы увидеть ситуацию яснее.\n\n"
        "Но если вы чувствуете, что хотите разобраться глубже, на сессии мы обычно делаем две вещи:\n"
        "разбираем, где именно застряла ситуация\n"
        "и находим точки выхода и ресурсы, на которые можно опереться.",
        reply_markup=InlineKeyboardMarkup(inline_keyboard=[
            [InlineKeyboardButton(text="Записаться на сессию", url=SESSION_URL)],
            [InlineKeyboardButton(text="🔄 Пройти ещё один разбор", callback_data="second")]
        ])
    )


@dp.callback_query(F.data == "second")
async def second_handler(callback: types.CallbackQuery, state: FSMContext):
    await callback.answer()
    await state.clear()
    await callback.message.answer(
        "Хорошо, начнём новый разбор.\n\nРасскажите — что на этот раз?",
        reply_markup=ReplyKeyboardRemove()
    )
    await state.set_state(Dialog.describe)


@dp.message(Command("stats"))
async def stats_handler(message: types.Message):
    if ADMIN_ID and message.from_user.id != ADMIN_ID:
        return
    try:
        stats = await get_stats()
    except Exception as e:
        await message.answer(f"Ошибка получения статистики: {e}")
        return

    total_sessions = stats["total_sessions"] or 0
    total_completed = stats["total_completed"] or 0
    conversion = round(total_completed / total_sessions * 100) if total_sessions else 0

    text = (
        f"📊 <b>Статистика бота</b>\n\n"
        f"👥 Уникальных пользователей: <b>{stats['total_users']}</b>\n"
        f"🔄 Всего сессий: <b>{total_sessions}</b>\n"
        f"✅ Дошли до разбора: <b>{total_completed}</b> ({conversion}%)\n"
        f"📅 Сегодня: <b>{stats['today_sessions']}</b>\n"
    )

    if stats["top_themes"]:
        text += f"\n🔥 <b>Популярные темы:</b>\n"
        for i, row in enumerate(stats["top_themes"], 1):
            text += f"  {i}. {row['theme']} — {row['cnt']}\n"

    if stats["feedback"]:
        fb = {r["value"]: r["cnt"] for r in stats["feedback"]}
        text += f"\n💬 <b>Отзывы ({sum(fb.values())}):</b>\n"
        text += f"  ✅ Да — {fb.get('да', 0)}\n"
        text += f"  🤔 Частично — {fb.get('частично', 0)}\n"
        text += f"  ❌ Нет — {fb.get('нет', 0)}\n"

    if stats["returning"]:
        text += f"\n🔁 <b>Возвращались:</b>\n"
        for row in stats["returning"]:
            name = row["username"] or str(row["user_id"])
            text += f"  @{name} — {row['sessions']} раз\n"

    await message.answer(text)

    if len(stats["all_themes"]) >= 5:
        analysis_prompt = (
            f"Вот список тем которые выбирали пользователи бота психологического разбора: {stats['all_themes']}. "
            f"Кратко (3-5 предложений): какие паттерны видны? Что чаще всего беспокоит аудиторию? "
            f"Какую тему стоит раскрыть в контенте коучу?"
        )
        response = anthropic_client.messages.create(
            model="claude-sonnet-4-6",
            max_tokens=400,
            messages=[{"role": "user", "content": analysis_prompt}]
        )
        await message.answer(f"🧠 <b>Анализ тем:</b>\n\n{response.content[0].text}")


@dp.message()
async def fallback_handler(message: types.Message, state: FSMContext):
    current_state = await state.get_state()
    if current_state is None:
        await message.answer(
            "Напишите /start чтобы начать разбор.",
            reply_markup=ReplyKeyboardRemove()
        )


async def send_reminders():
    while True:
        try:
            conn = await get_db()
            rows = await conn.fetch("""
                SELECT user_id, last_theme FROM users
                WHERE reminded = FALSE AND remind_at IS NOT NULL AND remind_at <= NOW()
            """)
            await conn.close()

            for row in rows:
                uid = row["user_id"]
                theme = row["last_theme"] or ""
                prompt = (
                    f"Пользователь 3 дня назад прошёл разбор ситуации на тему: «{theme}». "
                    f"Напиши короткое тёплое сообщение от Мадам Селезнёвой (2-3 предложения): "
                    f"спроси как они, упомни тему разбора, задай один мягкий вопрос о том, "
                    f"что изменилось или что удалось обдумать. "
                    f"Затем предложи новый разбор если появилась другая ситуация. "
                    f"Тон — тёплый, без навязчивости."
                )
                response = anthropic_client.messages.create(
                    model="claude-sonnet-4-6",
                    max_tokens=300,
                    messages=[{"role": "user", "content": prompt}]
                )
                reminder_text = response.content[0].text
                reminder_text += "\n\nЕсли хотите разобрать новую ситуацию — просто напишите /start"
                try:
                    await bot.send_message(uid, reminder_text)
                except Exception as e:
                    logging.error(f"Reminder send error for {uid}: {e}")
                conn2 = await get_db()
                await conn2.execute("UPDATE users SET reminded = TRUE WHERE user_id = $1", uid)
                await conn2.close()
        except Exception as e:
            logging.error(f"Reminder loop error: {e}")
        await asyncio.sleep(3600)


async def main():
    await init_db()
    asyncio.create_task(send_reminders())
    await dp.start_polling(bot)


if __name__ == "__main__":
    asyncio.run(main())
